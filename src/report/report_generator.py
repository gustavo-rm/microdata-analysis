from docx import Document

class ReportGenerator:
    def __init__(self, analysis_results, output_path):
        self.analysis_results = analysis_results
        self.output_path = output_path

    def generate_csv_report(self):
        """Gera um relatório em formato CSV."""
        self.analysis_results.to_csv(self.output_path + 'relatorio.csv', index=False)
        print("Relatório CSV gerado com sucesso!")

    def generate_word_report(self):
        """Gera um relatório em formato Word."""
        doc = Document()
        doc.add_heading('Análise de Dados do Mercado de Trabalho', level=1)
        doc.add_paragraph('Este relatório apresenta uma análise com base nos microdados fornecidos...')
        # Adiciona tabela com os resultados
        table = doc.add_table(rows=1, cols=len(self.analysis_results.columns))
        table.style = 'Table Grid'
        for i, col_name in enumerate(self.analysis_results.columns):
            table.cell(0, i).text = col_name
        for _, row in self.analysis_results.iterrows():
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)
        doc.save(self.output_path + 'relatorio.docx')
        print("Relatório Word gerado com sucesso!")

