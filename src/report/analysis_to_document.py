from docx import Document
from docx.shared import Pt
import os
from src.analysis.basic_statistics import BasicStatistics
from src.analysis.employment_indexes import EmploymentIndexes
from src.analysis.gender_analysis import GenderAnalysis
from src.analysis.position_analysis import PositionAnalysis
from src.analysis.regional_analysis import RegionalAnalysis
from src.analysis.predictive_models import PredictiveModels
from src.analysis.statistical_tests import StatisticalTests


class AnalysisToDocument:
    """
    Classe para executar métodos de análise e salvar os resultados em um arquivo Word.
    """

    def __init__(self, df, output_path="./data/processed/"):
        """
        Inicializa a classe com os dados e o caminho de saída.

        Parameters:
            df (pd.DataFrame): O conjunto de dados para análise.
            output_path (str): O caminho do arquivo Word onde os resultados serão salvos.
        """
        self.df = df
        self.output_path = output_path
        self.document = Document()

    def add_section(self, title):
        """
        Adiciona uma seção com título ao documento.

        Parameters:
            title (str): O título da seção.
        """
        self.document.add_heading(title, level=1)

    def add_subsection(self, subsection_title):
        """
        Adiciona um subtítulo ao documento.

        Parameters:
            subsection_title (str): O título do subtítulo a ser adicionado.
        """
        self.document.add_heading(subsection_title, level=2)

    def add_paragraph(self, text):
        """
        Adiciona um parágrafo ao documento com o texto fornecido.

        Parameters:
            text (str): O texto a ser adicionado ao parágrafo.
        """
        self.document.add_paragraph(text)

    def add_key_value_pair(self, key, value):
        """
        Adiciona um par chave/valor ao documento.

        Parameters:
            key (str): Nome do parâmetro.
            value (str): Valor associado ao parâmetro.
        """
        paragraph = self.document.add_paragraph()
        run_key = paragraph.add_run(f"{key}: ")
        run_key.bold = True
        paragraph.add_run(str(value))

    def add_table(self, data, column_names):
        """
        Adiciona uma tabela ao documento.

        Parameters:
            data (list ou pd.DataFrame): Os dados a serem adicionados na tabela.
                                         Pode ser uma lista de listas ou um DataFrame.
            column_names (list): Nomes das colunas da tabela.
        """
        if isinstance(data, list):
            # Caso os dados sejam uma lista de listas
            rows = data
        else:
            # Caso seja um DataFrame, converte para lista de listas
            rows = data.values.tolist()

        # Cria a tabela com o número correto de linhas e colunas
        table = self.document.add_table(rows=len(rows) + 1, cols=len(column_names))
        table.style = 'Table Grid'

        # Adiciona os nomes das colunas
        for idx, column_name in enumerate(column_names):
            table.cell(0, idx).text = column_name

        # Adiciona os dados
        for i, row in enumerate(rows):
            for j, value in enumerate(row):
                table.cell(i + 1, j).text = str(value)

    def save_to_word(self, file_name="data_analysis.docx"):
        """
        Salva o documento no arquivo especificado.

        Parameters:
            file_name (str, optional): Nome do arquivo de saída.
                                       Se não fornecido, será usado o nome padrão definido na inicialização.
        """
        # Garante que o caminho do diretório existe
        if not os.path.exists(self.output_path):
            os.makedirs(self.output_path)

        # Salva o documento no caminho especificado
        full_path = os.path.join(self.output_path, file_name)
        self.document.save(full_path)
        print(f"Documento salvo com sucesso em: {full_path}")

    def run_analysis(self):
        """
        Executa os métodos de análise e salva os resultados no arquivo Word.
        """
        self.add_section("Estatísticas Básicas")
        self.add_basic_statistics()

        self.add_section("Análise de Gênero")
        self.add_gender_analysis()

        self.add_section("Análise de Cargos")
        self.add_position_analysis()

        self.add_section("Previsões de Salário Médio para 2024")
        self.add_salary_predictions()

        self.add_section("Análise Regional")
        self.add_regional_analysis()

        self.add_section("Índices de Emprego")
        self.add_employment_indexes()

        self.add_section("Análise Estatística")
        self.add_statistical_tests()

        # Salvar o arquivo Word
        self.save_to_word()

    # Métodos Auxiliares
    def add_basic_statistics(self):
        """Adiciona estatísticas básicas ao relatório."""
        basic_stats = BasicStatistics(self.df)
        self.add_key_value_pair("Total de Empregados", basic_stats.calculate_total_employees())
        self.add_key_value_pair("Média Salarial Geral", f"{basic_stats.calculate_average_salary():,.2f}")

        self.add_subsection("Média Salarial por Ano")
        avg_salary_by_year = basic_stats.calculate_average_salary_by_year()
        self.add_table(avg_salary_by_year, column_names=["Ano", "Média Salarial"])

        self.add_subsection("Distribuição Salarial")
        salary_dist = basic_stats.salary_distribution()
        salary_dist_table = [
            ["Média Salarial", f"{salary_dist['Média Salarial']:,.2f}"],
            ["Mediana Salarial", f"{salary_dist['Mediana Salarial']:,.2f}"],
            ["Desvio Padrão", f"{salary_dist['Desvio Padrão']:,.2f}"],
            ["Coeficiente de Variação (%)", f"{salary_dist['Coeficiente de Variação (%)']:,.2f}"]
        ]
        self.add_table(salary_dist_table, column_names=["Métrica", "Valor"])

    def add_gender_analysis(self):
        """Adiciona a análise de gênero ao relatório."""
        gender_analysis = GenderAnalysis(self.df)

        self.add_subsection("Diferença Salarial entre Gêneros")
        salary_gap = gender_analysis.gender_salary_gap()
        for key, value in salary_gap.items():
            self.add_key_value_pair(key, f"{value:,.2f}")

        self.add_subsection("Métricas de Igualdade de Gênero")
        gender_equality = gender_analysis.gender_equality_metrics()
        for key, value in gender_equality.items():
            self.add_key_value_pair(key, f"{value:,.2f}" if value is not None else "N/A")

        self.add_subsection("Comparação Salarial por Gênero nos Top 10 Cargos")
        top_10_gender_salary = gender_analysis.compare_salary_by_gender_top_10_jobs()
        self.add_table(top_10_gender_salary, column_names=[
            "Cargo", "Salário Médio Masculino", "Salário Médio Feminino", "Diferença Salarial", "Total de Empregados"
        ])

        self.add_subsection("Empregados com Mais Vínculos Ativos por Ano")
        top_active_by_year = gender_analysis.top_active_employees_by_year(sigla_uf="PR")
        self.add_table(top_active_by_year, column_names=[
            "Ano", "Gênero", "ID Município", "Total de Vínculos Ativos"
        ])

    def add_position_analysis(self):
        """Adiciona a análise de cargos ao relatório."""
        position_analysis = PositionAnalysis(self.df)
        unique_positions = position_analysis.analyze_unique_positions(sort_by="Frequência", ascending=False)
        self.add_key_value_pair("Total de Cargos Diferentes", unique_positions["Total de Cargos Diferentes"])

        top_positions_df = unique_positions["Cargos"].head(10)
        self.add_table(top_positions_df, column_names=["Cargo", "Frequência"])

    def add_salary_predictions(self):
        """Adiciona as previsões de salário médio ao relatório."""
        predictive_models = PredictiveModels(self.df)
        gender_salary_predictions = predictive_models.predict_gender_salary_2024()

        for gender, predicted_salary in gender_salary_predictions.items():
            self.add_key_value_pair(f"Salário Médio Previsto ({gender})", f"{predicted_salary:,.2f}")

    def add_regional_analysis(self, average_by_city=False):
        """Adiciona a análise regional ao relatório."""
        regional_analysis = RegionalAnalysis(self.df)

        self.add_subsection("Total de Empregados e Salário Médio por Estado")
        regional_data = regional_analysis.regional_analysis()
        self.add_table(regional_data, column_names=["Estado", "Salário Médio", "Total Empregados"])

        self.add_subsection("Concentração de Empregos por Estado")
        concentration_state = regional_analysis.concentration_of_jobs(level="estado")
        self.add_table(concentration_state, column_names=["Estado", "Total de Empregados"])

        if average_by_city:
            self.add_subsection("Salário Médio por Município")
            average_salary_municipality = regional_analysis.average_salary_by_region(level="municipio")
            self.add_table(average_salary_municipality, column_names=["Município", "Salário Médio"])

        self.add_subsection("Média Salarial nas 5 Cidades Mais Populosas do Paraná")
        top_5_cities = regional_analysis.average_salary_top_5_cities()
        for city, avg_salary in top_5_cities.items():
            self.add_key_value_pair(city, f"{avg_salary:,.2f}" if avg_salary is not None else "N/A")

    def add_employment_indexes(self):
        """Adiciona índices de emprego ao relatório."""
        employment_indexes = EmploymentIndexes(self.df)
        self.add_key_value_pair("Índice de Disparidade Salarial (IDS)",
                                f"{employment_indexes.salary_disparity_index():,.2f}")
        self.add_key_value_pair("Índice de Escolaridade", f"{employment_indexes.education_index():,.2f}")

    def add_statistical_tests(self):
        """Adiciona os resultados de testes estatísticos ao relatório."""
        self.add_subsection("Teste ANOVA para Salários por Setor")
        self.add_paragraph(
            "O teste ANOVA (Análise de Variância) é utilizado para comparar as médias salariais "
            "entre diferentes setores. Um valor-p (p-value) menor que 0.05 indica que há uma diferença "
            "estatisticamente significativa entre os setores analisados."
        )

        statistical_tests = StatisticalTests(self.df)
        anova_results = statistical_tests.anova_salary_by_sector()

        if "Erro" in anova_results:
            self.add_paragraph(f"Erro: {anova_results['Erro']}")
        else:
            self.add_key_value_pair("Estatística F", f"{anova_results['Estatística']:.2f}")
            self.add_key_value_pair("Valor-p", f"{anova_results['Valor-p']:.4f}")

